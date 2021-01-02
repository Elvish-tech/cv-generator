from docx import Document
from docx.shared import Inches

document = Document()

# profile picture
document.add_picture('pc.jpg', width=Inches(2.0))

# name,ph.no,email
name = input('What is your name ?')
phone_number = input('What is your phone number ?')
email = input('What is your email ?')

document.add_paragraph(
    name + ' | ' + phone_number + ' | ' + email 
)
#about you
document.add_heading('About me')
document.add_paragraph(input('Tell about yourself'))

#work experience
document.add_heading('Work experience')
p=document.add_paragraph()

company = input('Enter company')
from_date = input('From date')
to_date = input('To date')

p.add_run(company + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

experience_details = input('Share your experience at ' + company)
p.add_run(experience_details)

#more experiences
while True:
    has_more_experiences = input('DO you have more experiences ? Yes/No')
    if has_more_experiences.lower() =='yes':
        p=document.add_paragraph()

        company = input('Enter company')
        from_date = input('From date')
        to_date = input('To date')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + '-' + to_date + '\n').italic = True

        experience_details = input('Share your experience at ' + company)
        p.add_run(experience_details)
    else:
        break
#skills
document.add_heading('Skills')
skills = input('Enter your skills ')
p = document.add_paragraph(skills)
p.style = 'List Bullet'

while True:
    has_more_skills = input('Do you have any additional skills ? Yes/No')
    if has_more_skills.lower() == 'yes':
        skills = input('Enter your skills')
        p =  document.add_paragraph(skills)
        p.style = 'List Bullet'
    else:
        break

#footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = 'CV generated using elvish-tech'


document.save('cv.docx')